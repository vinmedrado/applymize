"""CareerLens v3 - CV Exporter"""
import re
from pathlib import Path
from datetime import datetime

def generate_docx(text, contact, output_dir=None):
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        doc = Document()
        for s in doc.sections:
            s.top_margin = Cm(1.5); s.bottom_margin = Cm(1.5)
            s.left_margin = Cm(2.0); s.right_margin = Cm(2.0)
        HEADERS = ["RESUMO","OBJETIVO","EXPERIÊNCIA","EXPERIENCIA","FORMAÇÃO","FORMACAO",
                   "EDUCAÇÃO","EDUCACAO","HABILIDADES","COMPETÊNCIAS","PROJETOS",
                   "CERTIFICAÇÕES","CERTIFICACOES","IDIOMAS","CURSOS"]
        name = contact.get("name","Seu Nome")
        np = doc.add_paragraph(); np.alignment = WD_ALIGN_PARAGRAPH.CENTER
        nr = np.add_run(name.upper()); nr.bold = True; nr.font.size = Pt(18)
        nr.font.color.rgb = RGBColor(0x1A,0x1A,0x2E)
        parts = [p for p in [contact.get("email",""),contact.get("phone",""),
                              contact.get("linkedin",""),contact.get("github","")] if p]
        if parts:
            cp = doc.add_paragraph(" | ".join(parts)); cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in cp.runs: r.font.size = Pt(9); r.font.color.rgb = RGBColor(0x44,0x44,0x44)
        doc.add_paragraph("─"*85)
        for line in text.split('\n'):
            s = line.strip()
            if not s: continue
            is_h = any(s.upper().startswith(h) or s.upper()==h for h in HEADERS)
            if is_h:
                p = doc.add_paragraph(); r = p.add_run(s.upper())
                r.bold=True; r.font.size=Pt(11); r.font.color.rgb=RGBColor(0x16,0x3A,0x7C)
                d = doc.add_paragraph(); dr = d.add_run("─"*80)
                dr.font.size=Pt(6); dr.font.color.rgb=RGBColor(0x16,0x3A,0x7C)
            elif s.startswith(("•","-","*")):
                bp = doc.add_paragraph(style='List Bullet')
                br = bp.add_run(s.lstrip("•-* ").strip()); br.font.size=Pt(10)
            elif '|' in s or re.match(r'^\d{4}',s):
                p = doc.add_paragraph(); r = p.add_run(s); r.bold=True; r.font.size=Pt(10)
            else:
                p = doc.add_paragraph(); r = p.add_run(s); r.font.size=Pt(10)
        if output_dir is None: output_dir = Path.home()/".applymize"/"exports"
        Path(output_dir).mkdir(parents=True, exist_ok=True)
        ts = datetime.now().strftime("%Y%m%d_%H%M%S")
        path = Path(output_dir)/f"cv_applymize_{ts}.docx"
        doc.save(str(path)); return str(path)
    except Exception as e: return f"Erro: {e}"
