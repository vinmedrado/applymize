"""CareerLens v3 - CV Parser"""
import re
from pathlib import Path

def extract_text(path):
    s = Path(path).suffix.lower()
    try:
        if s == ".pdf":
            import pdfplumber
            t = ""
            with pdfplumber.open(path) as pdf:
                for pg in pdf.pages:
                    x = pg.extract_text()
                    if x: t += x + "\n"
            return t.strip()
        elif s in [".docx",".doc"]:
            from docx import Document
            doc = Document(path)
            parts = [p.text for p in doc.paragraphs if p.text.strip()]
            for tbl in doc.tables:
                for row in tbl.rows:
                    for cell in row.cells:
                        if cell.text.strip(): parts.append(cell.text.strip())
            return "\n".join(parts)
    except Exception as e:
        return f"Erro: {e}"
    return "Formato não suportado."

def extract_contact(text):
    info = {"email":"","phone":"","linkedin":"","github":"","name":""}
    m = re.search(r'[\w\.-]+@[\w\.-]+\.\w+', text)
    if m: info["email"] = m.group()
    m = re.search(r'(\+55\s?)?(\(?\d{2}\)?\s?)(\d{4,5}[-\s]?\d{4})', text)
    if m: info["phone"] = m.group().strip()
    m = re.search(r'linkedin\.com/in/[\w\-]+', text, re.I)
    if m: info["linkedin"] = m.group()
    m = re.search(r'github\.com/[\w\-]+', text, re.I)
    if m: info["github"] = m.group()
    lines = [l.strip() for l in text.split('\n') if l.strip()]
    if lines:
        f = lines[0]
        if len(f.split()) <= 6 and not any(c in f for c in ['@','/','.','|']):
            info["name"] = f
    return info

def parse_sections(text):
    kws = {
        "resumo":      ["resumo","objetivo","perfil","sobre mim"],
        "experiencia": ["experiência","experiencia","experience","histórico"],
        "educacao":    ["educação","educacao","formação","formacao","education"],
        "habilidades": ["habilidades","competências","skills","tecnologias"],
        "projetos":    ["projetos","projects","portfólio","portfolio"],
        "certificacoes":["certificações","certificados","certifications","cursos"],
        "idiomas":     ["idiomas","languages"],
    }
    sections = {k:"" for k in kws}
    current = None
    for line in text.split('\n'):
        ll = line.lower().strip()
        hit = False
        for sec, keys in kws.items():
            if any(k in ll for k in keys) and len(line.strip()) < 50:
                current = sec; hit = True; break
        if not hit and current:
            sections[current] += line + "\n"
    return {k:v.strip() for k,v in sections.items() if v.strip()}

def full_parse(path):
    text = extract_text(path)
    if text.startswith("Erro"): return {"error": text, "raw_text": ""}
    return {
        "raw_text": text,
        "contact": extract_contact(text),
        "sections": parse_sections(text),
        "word_count": len(text.split()),
    }
